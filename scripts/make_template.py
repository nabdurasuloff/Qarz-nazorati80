# -*- coding: utf-8 -*-
"""
Ogohlantirish xati / Talabnoma uchun Word shablonini yaratadi.
Shablonda {{PLACEHOLDER}} ko'rinishidagi joylar keyinchalik dastur
tomonidan haqiqiy ma'lumotlar bilan almashtiriladi.
"""
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para(doc, text='', size=11, bold=False, italic=False, align=None,
             space_after=8, font_name='Times New Roman', color=None,
             indent_first=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent_first is not None:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def build():
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # ---- Header: bank rekvizitlari (chapda) ----
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Cm(11)
    header_table.columns[1].width = Cm(6.5)
    left_cell, right_cell = header_table.rows[0].cells
    set_cell_border_none(left_cell)
    set_cell_border_none(right_cell)

    left_lines = [
        "{{BANK_MANZIL}}",
        "{{BANK_EMAIL}}",
        "{{BANK_SAYT}}",
        "Tel: {{BANK_TEL}}",
    ]
    lp = left_cell.paragraphs[0]
    lp.text = ''
    for i, line in enumerate(left_lines):
        r = lp.add_run(('\n' if i else '') + line)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x2F, 0x6B, 0x8A)

    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = rp.add_run("{{BANK_NOMI}}")
    r.bold = True
    r.font.size = Pt(11)

    add_para(doc, '', size=6, space_after=4)

    # ---- Qabul qiluvchi ----
    add_para(doc, "{{MIJOZ_ISM}}", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "({{MIJOZ_MANZIL}})", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    # ---- Sarlavha ----
    add_para(doc, "{{SARLAVHA}}", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # ---- Asosiy matn ----
    add_para(
        doc,
        '"{{BANK_QISQA_NOMI}}" АТБ {{FILIAL_NOMI}} филиали Сизга шуни маълум қиладики, '
        'Банк ва Сизнинг ўртангизда имзоланган кредит шартномасига асосан '
        '{{KREDIT_SUMMA}} сўм миқдорида {{KREDIT_MAQSAD}} кредити ажратилган.',
        size=11.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=12
    )

    add_para(
        doc,
        '"{{BANK_QISQA_NOMI}}" АТБ {{FILIAL_NOMI}} филиали томонидан Сизга ажратилган кредит '
        'бўйича {{HOLAT_SANASI}} холатига жами муддати ўтган {{JAMI_QARZ}} сўм қарздорлик мавжуд '
        'булиб, бунда муддати ўтган асосий қарздорлик {{ASOSIY_QARZ}} сўмни, муддати ўтган фоиз '
        'қарздорлик {{FOIZ_QARZ}} сўм хамда кредит тўлови ўз вақтидан кечиктирилганлиги учун '
        'хисобланган жарима (пеня) {{JARIMA}} сўмни ташкил қилади.',
        size=11.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=12
    )

    add_para(
        doc,
        '"{{BANK_QISQA_NOMI}}" АТБ {{FILIAL_NOMI}} филиали, ажратилган кредит бўйича шартномавий '
        'мажбуриятларни лозим даражада бажаришни таъминлаган холда банк олдидаги муддати ўтган '
        'қарздорликни {{TOLOV_MUDDATI}} ичида тўлиқ тулашингизни талаб қилади.',
        size=11.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=12
    )

    add_para(
        doc,
        'Шунингдек, белгиланган муддат ичида муддати ўтган қарздорлик тўланмаган тақдирда кредит '
        'маблағини муддатидан олдин тўлиқ ундириш юзасидан тегишлилиги бўйича Суд органларига '
        'мурожаат қилиниб, қарздор, кафил ва биргаликда қарз олувчиларнинг иш ҳақидан ундириш ёки '
        'гаров таъминоти сифатида тақдим қилинган мол-мулкларга қаратиш бўйича даъво аризаси '
        'киритилади. Бунда, Суд органларига мурожаат қилиш, ишни судда кўриб чиқиш билан боғлиқ '
        'барча харажатлар, хусусан давлат божи, ижро йиғимлари ва жарималар Сизнинг зиммангизга '
        'юклатилиши ҳақида Сизни огоҳлантирамиз.',
        size=11.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=14
    )

    add_para(
        doc,
        'Эслатма: (қарздорлик суммасини Банкнинг исталган таркибий бўлинмаси (кассаси) орқали '
        'нақд пул шаклида ёки банк тўлов картаси орқали, шунингдек Банкнинг "{{BANK_MOBIL_ILOVA}}" '
        'мобил иловаси ёрдамида (кредит анкета № {{ANKETA_RAQAM}}, банк коди - {{BANK_KODI}}) '
        'банк филиалига ташриф буюрмасдан туриб тўлашингиз мумкинлигини маълум қиламиз. Мобил илова '
        'орқали тўловни амалга оширишда Сиз томонингиздан саволлар юзага келса Банкнинг алоқа '
        'марказига мурожаат қилишингиз мумкин. (тел: {{ALOQA_MARKAZI_TEL}})',
        size=10.5, italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=28
    )

    # ---- Imzo qismi ----
    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.columns[0].width = Cm(11)
    sign_table.columns[1].width = Cm(6.5)
    lcell, rcell = sign_table.rows[0].cells
    set_cell_border_none(lcell)
    set_cell_border_none(rcell)

    lp2 = lcell.paragraphs[0]
    r1 = lp2.add_run('"{{BANK_QISQA_NOMI}}" АТБ\n{{FILIAL_NOMI}} филиали бошқарувчиси\n\n')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = lp2.add_run('Мурожаат учун тел: {{FILIAL_TEL}}')
    r2.italic = True
    r2.font.size = Pt(10.5)

    rp2 = rcell.paragraphs[0]
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r3 = rp2.add_run('\n{{RAHBAR_ISM}}')
    r3.bold = True
    r3.font.size = Pt(11)

    doc.save('/home/claude/qarz_nazorat/templates/xat_shablon.docx')
    print("Shablon yaratildi.")


if __name__ == '__main__':
    build()
