# -*- coding: utf-8 -*-
"""
O'zbek kirill alifbosidan lotin alifbosiga o'girish (transliteratsiya).
Faqat shablon fayllarini bir martalik yangilash uchun ishlatiladi.
"""
import re

# Ko'p harfli birikmalar avval almashtirilishi kerak (uzunroq ketma-ketlik ustuvor)
MULTI = [
    ('Ё', 'Yo'), ('ё', 'yo'),
    ('Ю', 'Yu'), ('ю', 'yu'),
    ('Я', 'Ya'), ('я', 'ya'),
    ('Ц', 'Ts'), ('ц', 'ts'),
    ('Ч', 'Ch'), ('ч', 'ch'),
    ('Ш', 'Sh'), ('ш', 'sh'),
    ('Щ', 'Sh'), ('щ', 'sh'),
    ('Ғ', "G'"), ('ғ', "g'"),
    ('Ў', "O'"), ('ў', "o'"),
    ('Ҳ', 'H'), ('ҳ', 'h'),
    ('Нг', 'Ng'), ('нг', 'ng'),
]

SINGLE = {
    'А': 'A', 'а': 'a', 'Б': 'B', 'б': 'b', 'В': 'V', 'в': 'v',
    'Г': 'G', 'г': 'g', 'Д': 'D', 'д': 'd', 'Ж': 'J', 'ж': 'j',
    'З': 'Z', 'з': 'z', 'И': 'I', 'и': 'i', 'Й': 'Y', 'й': 'y',
    'К': 'K', 'к': 'k', 'Қ': 'Q', 'қ': 'q', 'Л': 'L', 'л': 'l',
    'М': 'M', 'м': 'm', 'Н': 'N', 'н': 'n', 'О': 'O', 'о': 'o',
    'П': 'P', 'п': 'p', 'Р': 'R', 'р': 'r', 'С': 'S', 'с': 's',
    'Т': 'T', 'т': 't', 'У': 'U', 'у': 'u', 'Ф': 'F', 'ф': 'f',
    'Х': 'X', 'х': 'x', 'Э': 'E', 'э': 'e', 'Ы': 'I', 'ы': 'i',
    'Ъ': "'", 'ъ': "'", 'Ь': '', 'ь': '',
}


def _e_replace(text):
    """'Е'/'е' so'z boshida yoki unlidan keyin 'Ye'/'ye', aks holda 'E'/'e'."""
    def repl(m):
        prefix = m.group(1) or ''
        ch = m.group(2)
        is_start = (prefix == '' or prefix in ' \t\n(«"\'—-.,;:!?')
        if is_start:
            return prefix + ('Ye' if ch == 'Е' else 'ye')
        return prefix + ('E' if ch == 'Е' else 'e')
    return re.sub(r'(.?)([Ее])', repl, text)


def cyr_to_lat(text):
    """Berilgan matndagi kirill (o'zbekcha) matnni lotin alifbosiga o'giradi.
    {{PLACEHOLDER}} ko'rinishidagi joylar (allaqachon lotin/UPPERCASE) o'zgarishsiz qoladi."""
    # Placeholderlarni vaqtincha ajratib olamiz, ular ichida kirill harf yo'q
    placeholders = []

    def stash(m):
        placeholders.append(m.group(0))
        return f"\x00{len(placeholders) - 1}\x00"

    text = re.sub(r'\{\{[^}]+\}\}', stash, text)

    text = _e_replace(text)
    for cyr, lat in MULTI:
        text = text.replace(cyr, lat)
    for cyr, lat in SINGLE.items():
        text = text.replace(cyr, lat)

    def unstash(m):
        return placeholders[int(m.group(1))]

    text = re.sub(r'\x00(\d+)\x00', unstash, text)
    return text


if __name__ == '__main__':
    from docx import Document

    src = '/home/claude/qarz_nazorat/templates/xat_shablon.docx'
    doc = Document(src)

    def convert_runs(paragraph):
        for run in paragraph.runs:
            if run.text:
                run.text = cyr_to_lat(run.text)

    for p in doc.paragraphs:
        convert_runs(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    convert_runs(p)

    doc.save(src)
    print("Shablon lotin alifbosiga o'girildi:", src)
