# -*- coding: utf-8 -*-
import re
from docx import Document


def _replace_in_paragraph_regex(paragraph, patterns):
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text:
        return
    new_text = full_text
    for pattern, repl in patterns:
        new_text = re.sub(pattern, repl, new_text)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def templatize(input_path, output_path, patterns):
    doc = Document(input_path)
    for p in doc.paragraphs:
        _replace_in_paragraph_regex(p, patterns)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_regex(p, patterns)
    doc.save(output_path)
    return output_path


# =====================================================================
# 1) Jismoniy shaxsga oddiy da'vo ariza (kafilsiz)
# =====================================================================
PATTERNS_JISMONIY_ODDIY = [
    (r'2026 йил.*?-сонли\s*', '{{ARIZA_SANA_QATORI}}\n'),
    (r'Фуқаролик ишлари бўйича.*?судига', '{{SUD_NOMI}}'),
    (r'Ундурувчи:\t.*АТБ\s*', 'Ундурувчи:\t{{BANK_NOMI}} '),
    (r'ҳ\|р:\s*[\d]+.*?уй', '{{BANK_REKVIZIT_QATORI}}'),
    (r'Қарздор:\t.*?,\s*', 'Қарздор:\t{{MIJOZ_ISM}}, '),
    (r'Манзили:.*МФЙ', 'Манзили:{{MIJOZ_MANZIL}}'),
    (r'ПИНФЛ:\s*\d+', 'ПИНФЛ: {{MIJOZ_PINFL}}'),
    (r'\(паспорт маълумоти:.*?берилган\)', '(паспорт маълумоти: {{MIJOZ_PASSPORT_TOLIQ}})'),
    (r'“Агробанк” АТБ Сирдарё вилоят худудий филиали билан қарз олувчи .*? ўртасида тузилган .*? йилдаги кредит шартномасига биноан банк томонидан қарз олувчига йиллик .*? фоиз устама ҳақ тўлаш шарти билан .*? ой муддатга .*? сўм миқдорида кредит маблағи ажратилган\.',
     '"{{BANK_QISQA_NOMI}}" АТБ {{BANK_HUDUDIY_NOMI}} билан қарз олувчи {{MIJOZ_ISM}} ўртасида тузилган '
     '{{SHARTNOMA_SANA}} йилдаги кредит шартномасига биноан банк томонидан қарз олувчига йиллик '
     '{{YILLIK_FOIZ}} фоиз устама ҳақ тўлаш шарти билан {{KREDIT_MUDDATI_OY}} ой муддатга '
     '{{KREDIT_SUMMA}} сўм миқдорида кредит маблағи ажратилган.'),
    (r'Натижада .*? ҳолатига қарз олувчининг банк олдида жами .*? сўм \(Бунда, муддати ўтган қарздорлик .*? сўм, фоиз қарздорлик .*? сўм, ҳамда пеня .*? сўм\), қарздорлиги мавжуд\.',
     'Натижада {{HOLAT_SANASI}} ҳолатига қарз олувчининг банк олдида жами {{JAMI_QARZ}} сўм '
     '(Бунда, муддати ўтган қарздорлик {{ASOSIY_QARZ}} сўм, фоиз қарздорлик {{FOIZ_QARZ}} сўм, '
     'ҳамда пеня {{JARIMA}} сўм), қарздорлиги мавжуд.'),
    (r'Қарздор .*?дан\s+солидар тартибда “Агробанк” АТБ Сирдарё вилоят худудий бошкармаси \(Боёвут филиали\) фойдасига жами .*? сўм \(Бунда, муддати ўтган қарздорлик .*? сўм, фоиз қарздорлик .*? сўм, ҳамда пеня .*? сўм\) сўм кредит қарздорлигини ҳамда олдиндан тўланган почта ҳаражатларини ундиришни;',
     'Қарздор {{MIJOZ_ISM}}дан "{{BANK_QISQA_NOMI}}" АТБ {{BANK_HUDUDIY_NOMI}} фойдасига жами '
     '{{JAMI_QARZ}} сўм (Бунда, муддати ўтган қарздорлик {{ASOSIY_QARZ}} сўм, фоиз қарздорлик '
     '{{FOIZ_QARZ}} сўм, ҳамда пеня {{JARIMA}} сўм) кредит қарздорлигини ҳамда олдиндан тўланган '
     '{{POCHTA_XARAJATI}} сўм почта ҳаражатларини ундиришни;'),
    (r'С\.Р\.Амонлиқов', '{{IMZO_ISM}}'),
    (r'Бошқарма бошлиғи ўринбосари:', '{{IMZO_LAVOZIM}}:'),
]

# =====================================================================
# 2) Jismoniy shaxsga da'vo + kafil (kafillikka qaratish)
# =====================================================================
PATTERNS_JISMONIY_KAFIL = [
    (r'2026 йил.*?-сонли\s*', '{{ARIZA_SANA_QATORI}}\n'),
    (r'Фуқаролик ишлари бўйича.*?судига', '{{SUD_NOMI}}'),
    (r'Ундурувчи:\t.*АТБ\s*', 'Ундурувчи:\t{{BANK_NOMI}} '),
    (r'ҳ\|р:\s*[\d]+.*?уй', '{{BANK_REKVIZIT_QATORI}}'),
    (r'ПИНФЛ:\s*32605952870012', 'ПИНФЛ: {{MIJOZ_PINFL}}'),
    (r'\(паспорт маълумоти: AE4339166.*?берилган\)', '(паспорт маълумоти: {{MIJOZ_PASSPORT_TOLIQ}})'),
    (r'ПИНФЛ:\s*32409922870028', 'ПИНФЛ: {{KAFIL_PINFL}}'),
    (r'\(паспорт маълумоти: AA9459035.*?берилган\)', '(паспорт маълумоти: {{KAFIL_PASSPORT_TOLIQ}})'),
    (r'“Агробанк” АТБ Сирдарё вилоят Боёвут филиали билан қарз олувчи .*? ўртасида тузилган .*? йилдаги кредит шартномасига биноан банк томонидан қарз олувчига йиллик .*? фоиз устама ҳақ тўлаш шарти билан .*? ой муддатга .*? сўм миқдорида кредит маблағи ажратилган\.',
     '"{{BANK_QISQA_NOMI}}" АТБ {{BANK_HUDUDIY_NOMI}} билан қарз олувчи {{MIJOZ_ISM}} ўртасида тузилган '
     '{{SHARTNOMA_SANA}} йилдаги кредит шартномасига биноан банк томонидан қарз олувчига йиллик '
     '{{YILLIK_FOIZ}} фоиз устама ҳақ тўлаш шарти билан {{KREDIT_MUDDATI_OY}} ой муддатга '
     '{{KREDIT_SUMMA}} сўм миқдорида кредит маблағи ажратилган.'),
    (r'Кредит таьминоти сифатида .*?нинг кафиллик шартномаси тақдим этилган\.',
     'Кредит таьминоти сифатида {{KAFIL_ISM}}нинг кафиллик шартномаси тақдим этилган.'),
    (r'Натижада .*? ҳолатига қарз олувчининг банк олдида жами .*? сўм \(Бунда, муддати ўтган қарздорлик .*? сўм, фоиз қарздорлик .*? сўм, ҳамда пеня .*? сўм\), қарздорлиги мавжуд\.',
     'Натижада {{HOLAT_SANASI}} ҳолатига қарз олувчининг банк олдида жами {{JAMI_QARZ}} сўм '
     '(Бунда, муддати ўтган қарздорлик {{ASOSIY_QARZ}} сўм, фоиз қарздорлик {{FOIZ_QARZ}} сўм, '
     'ҳамда пеня {{JARIMA}} сўм), қарздорлиги мавжуд.'),
    (r'Қарздор .*? ва .*?лардан\s+солидар тартибда “Агробанк” АТБ Сирдарё вилоят худудий бошкармаси \(Боёвут филиали\) фойдасига жами .*? сўм \(Бунда, муддати ўтган қарздорлик .*? сўм, фоиз қарздорлик .*? сўм, ҳамда пеня .*? сўм\) сўм кредит қарздорлигини ҳамда олдиндан тўланган почта ҳаражатларини ундиришни;',
     'Қарздор {{MIJOZ_ISM}} ва {{KAFIL_ISM}}дан солидар тартибда "{{BANK_QISQA_NOMI}}" АТБ '
     '{{BANK_HUDUDIY_NOMI}} фойдасига жами {{JAMI_QARZ}} сўм (Бунда, муддати ўтган қарздорлик '
     '{{ASOSIY_QARZ}} сўм, фоиз қарздорлик {{FOIZ_QARZ}} сўм, ҳамда пеня {{JARIMA}} сўм) кредит '
     'қарздорлигини ҳамда олдиндан тўланган {{POCHTA_XARAJATI}} сўм почта ҳаражатларини ундиришни;'),
    (r'С\.Р\.Амонлиқов', '{{IMZO_ISM}}'),
    (r'Бошқарма бошлиғи ўринбосари:', '{{IMZO_LAVOZIM}}:'),
]

JISMONIY_KAFIL_INDEX_FIX = {
    11: '  Қарздор:\t{{MIJOZ_ISM}}, ',
    12: 'Манзили:{{MIJOZ_MANZIL}}',
    17: '  жавобгар:\t{{KAFIL_ISM}}, ',
    18: 'Манзили:{{KAFIL_MANZIL}}',
}

# =====================================================================
# 3) Yuridik/Ф.Х iqtisodiy sudga (kafilsiz, Palataga to'g'ridan-to'g'ri)
# =====================================================================
PATTERNS_YURIDIK_ODDIY = [
    (r'2026 йил.*?-сонли', '{{ARIZA_SANA_QATORI}}'),
    (r'^Гулистон туманлараро иқтисодий судига$', '{{SUD_NOMI}}'),
    (r'^Ўзбекистон\s+Савдо-саноат Палатаси Сирдарё вилояти ҳудудий бошқармаси\s*$', '{{PALATA_NOMI}}'),
    (r'^“Агробанк” АТБ$', '{{BANK_NOMI}}'),
    (r'^Банк коди: \d+,$', 'Банк коди: {{BANK_KODI_BOSH}},'),
    (r'^СТИР: \d+,\s*$', 'СТИР: {{BANK_STIR}},'),
    (r'^Х\|р:\d+\.$', 'Х|р:{{BANK_HISOB_RAQAM_BOSH}}.'),
    (r'^Манзил: Тошкент шахри,Муқумий кўчаси 43-уй$', 'Манзил: {{BANK_MANZIL_BOSH}}'),
    (r'^«АЛЛАЁР-ОТА-АВЛОДЛАРИ Ф/Х» $', '{{MIJOZ_ISM}} '),
    (r'^Манзил;.*?маҳалласи\s*$', 'Манзил;{{MIJOZ_MANZIL}}'),
    (r'^Банк коди:\d+,$', 'Банк коди:{{MIJOZ_BANK_KODI}},'),
    (r'^Раҳбари:.*$', 'Раҳбари:{{MIJOZ_RAHBAR}}'),
    (r'^Телфон рақами; \+\d+$', 'Телфон рақами; {{MIJOZ_TEL}}'),
    (r'Ўзбекистон Савдо-саноат палатаси аъзоси ҳисобланган “Агробанк” АТБ \(Боёвут филиали\) Палатанинг Сирдарё вилоят ҳудудий бошқармасига мурожаат қилиб, унда банк манфаатини ҳимоя қилиб «АЛЛАЁР-ОТА-АВЛОДЛАРИ Ф/Х»нинг муддати ўтган кредит асосий қарзи, фоиз қарзи ҳамда кечиктирилган кунлар учун пеня ундириш ҳақида судга даъво аризаси киритиб берилишини сўраган\.',
     'Ўзбекистон Савдо-саноат палатаси аъзоси ҳисобланган "{{BANK_QISQA_NOMI}}" АТБ ({{FILIAL_NOMI}} филиали) '
     'Палатанинг {{PALATA_NOMI}}сига мурожаат қилиб, унда банк манфаатини ҳимоя қилиб {{MIJOZ_ISM}}нинг '
     'муддати ўтган кредит асосий қарзи, фоиз қарзи ҳамда кечиктирилган кунлар учун пеня ундириш ҳақида '
     'судга даъво аризаси киритиб берилишини сўраган.'),
    (r'Аниқланишича, «Агробанк» АТБ Боёвут филиали \(бундан кейин матнда “Банк” деб юритилади\) ва «АЛЛАЁР-ОТА-АВЛОДЛАРИ Ф/Х» \(бундан кейин матнда “Қарздор” деб юритилади\) ўртасида .*? йилда кредит шартномаси тузилган бўлиб, унга биноан банк томонидан қарздорга йиллик .*? фоиз устама ҳақ тўлаш шарти билан, “.*?” учун\s+.*? ой муддат билан қайтариш шарти билан .*? сўм миқдорида кредит маблағи ажратилиши белгиланган\.',
     'Аниқланишича, "{{BANK_QISQA_NOMI}}" АТБ {{FILIAL_NOMI}} филиали (бундан кейин матнда "Банк" деб '
     'юритилади) ва {{MIJOZ_ISM}} (бундан кейин матнда "Қарздор" деб юритилади) ўртасида '
     '{{SHARTNOMA_SANA}} йилда кредит шартномаси тузилган бўлиб, унга биноан банк томонидан қарздорга '
     'йиллик {{YILLIK_FOIZ}} фоиз устама ҳақ тўлаш шарти билан, "{{KREDIT_MAQSAD}}" учун '
     '{{KREDIT_MUDDATI_OY}} ой муддат билан қайтариш шарти билан {{KREDIT_SUMMA}} сўм миқдорида '
     'кредит маблағи ажратилиши белгиланган.'),
    (r'Ҳусусан, банк томонидан тақдим этилган маълумотномага кўра, қарздорнинг .*? йил ҳолатига банк олдида жами муддати ўтган .*? \(муддати ўтган асосий қарзи .*? сўм, .*? сўм фоиз қарзи, график режа асосида тўловни амалга оширмаганлиги учун ҳисобланган пеня .*?\) сўм кредит қарздорлиги юзага келган\.',
     'Ҳусусан, банк томонидан тақдим этилган маълумотномага кўра, қарздорнинг {{HOLAT_SANASI}} йил '
     'ҳолатига банк олдида жами муддати ўтган {{JAMI_QARZ}} (муддати ўтган асосий қарзи {{ASOSIY_QARZ}} '
     'сўм, {{FOIZ_QARZ}} сўм фоиз қарзи, график режа асосида тўловни амалга оширмаганлиги учун '
     'ҳисобланган пеня {{JARIMA}}) сўм кредит қарздорлиги юзага келган.'),
    (r'Жавобгар «АЛЛАЁР-ОТА-АВЛОДЛАРИ Ф/Х» ҳисобидан “Агробанк” АТБ фойдасига жами .*? \(муддати ўтган асосий қарзи .*? сўм, .*? сўм фоиз қарзи, график режа асосида тўловни амалга оширмаганлиги учун ҳисобланган пеня .*?\) сўм миқдоридаги кредит қарздорлигини хамда .*? сўм почта харажатини ундириш тўғрисида ҳал қилув қарори қабул қилишингизни;',
     'Жавобгар {{MIJOZ_ISM}} ҳисобидан "{{BANK_QISQA_NOMI}}" АТБ фойдасига жами {{JAMI_QARZ}} '
     '(муддати ўтган асосий қарзи {{ASOSIY_QARZ}} сўм, {{FOIZ_QARZ}} сўм фоиз қарзи, график режа '
     'асосида тўловни амалга оширмаганлиги учун ҳисобланган пеня {{JARIMA}}) сўм миқдоридаги кредит '
     'қарздорлигини хамда {{POCHTA_XARAJATI}} сўм почта харажатини ундириш тўғрисида ҳал қилув '
     'қарори қабул қилишингизни;'),
    (r'Сирдарё вилояти, Боёвут тумани, Боёвут шахарчаси, Тинчлик кўчаси, 10-уй Боёвут филиалига',
     '{{BANK_MANZIL_FILIAL}} {{FILIAL_NOMI}} филиалига'),
    (r'С\.Р\.Амонликов', '{{IMZO_ISM}}'),
    (r'Бошқарма бошлиғи\s*\n?\s*ўринбосари\s*:', '{{IMZO_LAVOZIM}}:'),
]

# =====================================================================
# 4) Yuridik shaxsga da'vo + kafil, foiz-penya undirish (Palataga)
# =====================================================================
PATTERNS_YURIDIK_KAFIL = [
    (r"O‘zbekiston Savdo-sanoat palatasi a’zosi hisoblangan \"Agrobank\" ATB Palataning Sirdaryo viloyat hududiy boshqarmasiga murojaat qilib, unda bank manfaatini himoya qilib «.*?» MChJ va qo‘shimcha javobgar .*?dan solidar tartibda muddati o‘tgan asosiy kredit va unga hisoblangan foiz hamda penya qarzdorligini undirish haqida sudga da’vo arizasi kiritib berilishini so‘ragan\.",
     "O'zbekiston Savdo-sanoat palatasi a'zosi hisoblangan \"{{BANK_QISQA_NOMI}}\" ATB Palataning "
     "{{PALATA_NOMI}}siga murojaat qilib, unda bank manfaatini himoya qilib {{MIJOZ_ISM}} va qo'shimcha "
     "javobgar {{KAFIL_ISM}}dan solidar tartibda muddati o'tgan asosiy kredit va unga hisoblangan foiz "
     "hamda penya qarzdorligini undirish haqida sudga da'vo arizasi kiritib berilishini so'ragan."),
    (r'Aniqlanishicha, «Agrobank» ATB Boyovut filiali \(bundan keyin matnda "Bank" deb yuritiladi\) va «.*?» MChJ \(bundan keyin matnda "Qarzdor" deb yuritiladi \) o‘rtasida 29\.12\.2023 yildagi №1 sonli   kredit shartnomasi tuzilgan bo‘lib,unga binoan bank tomonidan qarizdarga yillik 26 foiz ustama xaы tщlash sharti bilan, .*?sotib olish uchun 3 \(uch\) yil muddatga 2 380 000 000 \(Ikki milliard uch yuz sakson million\) so‘m miqdorida naqd pulsiz shaklda, pul ko‘chirish yo‘li bilan kredit mablag‘i ajratilgan\.',
     'Aniqlanishicha, "{{BANK_QISQA_NOMI}}" ATB {{FILIAL_NOMI}} filiali (bundan keyin matnda "Bank" deb '
     'yuritiladi) va {{MIJOZ_ISM}} (bundan keyin matnda "Qarzdor" deb yuritiladi) o\'rtasida '
     '{{SHARTNOMA_SANA}} yildagi kredit shartnomasi tuzilgan bo\'lib, unga binoan bank tomonidan '
     'qarzdorga yillik {{YILLIK_FOIZ}} foiz ustama haqi to\'lash sharti bilan, "{{KREDIT_MAQSAD}}" uchun '
     '{{KREDIT_MUDDATI_OY}} oy muddatga {{KREDIT_SUMMA}} so\'m miqdorida kredit mablag\'i ajratilgan.'),
    (r"Ajratilgan kredit mablag‘ining ta’minoti sifatida .*? yildagi №.*? sonli kafillik shartnomasiga asosan .*? fuqaro .*?ning kafilligi olingan\.",
     "Ajratilgan kredit mablag'ining ta'minoti sifatida {{SHARTNOMA_SANA}} yildagi kafillik shartnomasiga "
     "asosan {{KAFIL_MANZIL}} manzilida yashovchi fuqaro {{KAFIL_ISM}}ning kafilligi olingan."),
    (r'qarzdorning .*? yil holatiga bank oldida jami .*? so‘m \(shundan  .*? asosiy qarz, .*? so‘m foiz qarzi hamda penya .*? so‘m\) kredit qarzdorligi yuzaga kelgan\.',
     "qarzdorning {{HOLAT_SANASI}} yil holatiga bank oldida jami {{JAMI_QARZ}} so'm (shundan "
     "{{ASOSIY_QARZ}} asosiy qarz, {{FOIZ_QARZ}} so'm foiz qarzi hamda penya {{JARIMA}} so'm) kredit "
     "qarzdorligi yuzaga kelgan."),
    (r'1\.«.*?» MChJ va qo‘shimcha javobgar .*?dan solidar tartibda "Agrobank" ATB foydasiga jami .*? so‘m, \(shundan  .*? asosiy qarz, .*? so‘m foiz qarzi hamda penya .*? so‘m\) kredit qarzdorligi yuzaga kelgan kredit qarzdorligini undirish xamda .*? so‘m pochta xarajatini salidar tartibda undirishingizni;',
     "1.{{MIJOZ_ISM}} va qo'shimcha javobgar {{KAFIL_ISM}}dan solidar tartibda \"{{BANK_QISQA_NOMI}}\" ATB "
     "foydasiga jami {{JAMI_QARZ}} so'm (shundan {{ASOSIY_QARZ}} asosiy qarz, {{FOIZ_QARZ}} so'm foiz "
     "qarzi hamda penya {{JARIMA}} so'm) kredit qarzdorligini undirish xamda {{POCHTA_XARAJATI}} so'm "
     "pochta xarajatini salidar tartibda undirishingizni;"),
    (r'Sirdaryo viloyati, Boyovut tumani, Boyovut shaxarchasi, Tinchlik ko‘chasi, 10-uy Boyovut filialiga',
     '{{BANK_MANZIL_FILIAL}} {{FILIAL_NOMI}} filialiga'),
    (r'S\.R\.Amanliqov', '{{IMZO_ISM}}'),
    (r"Boshqarma boshlig‘i\s*\no‘rinbosari\s*:", "{{IMZO_LAVOZIM}}:"),
]

def set_table_cell_paragraph(doc, table_idx, row_idx, col_idx, para_idx, new_text):
    cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
    p = cell.paragraphs[para_idx]
    if p.runs:
        p.runs[0].text = new_text
        for run in p.runs[1:]:
            run.text = ''
    else:
        p.add_run(new_text)



def templatize_by_index(input_path, output_path, index_replacements, patterns=None):
    """
    index_replacements: dict {paragraph_index: new_full_text}
    Ambiguous (bir xil matn ikki joyda takrorlangan) qismlar uchun ishlatiladi.
    """
    doc = Document(input_path)
    for idx, new_text in index_replacements.items():
        p = doc.paragraphs[idx]
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ''
        else:
            p.add_run(new_text)
    if patterns:
        for p in doc.paragraphs:
            _replace_in_paragraph_regex(p, patterns)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph_regex(p, patterns)
    doc.save(output_path)
    return output_path


YURIDIK_KAFIL_TABLE_FIX = [
    # (table_idx, row_idx, col_idx, para_idx, new_text)
    (1, 0, 2, 2, '{{SUD_NOMI}} '),
    (1, 1, 2, 0, '{{PALATA_NOMI}} '),
    (1, 2, 2, 0, '"{{BANK_QISQA_NOMI}}" ATB manfaatini himoya qilib '),
    (1, 3, 2, 0, '{{BANK_NOMI}}'),
    (1, 4, 2, 0, 'Bank kodi: {{BANK_KODI_BOSH}}, '),
    (1, 5, 2, 0, 'STIR: {{BANK_STIR}}, '),
    (1, 6, 2, 0, 'X|r:{{BANK_HISOB_RAQAM_BOSH}}.'),
    (1, 7, 2, 0, 'Manzil: {{BANK_MANZIL_BOSH}}'),
    (1, 8, 2, 0, '\xa0{{MIJOZ_ISM}}'),
    (1, 9, 2, 0, '{{MIJOZ_MANZIL}}  Bank kodi:{{MIJOZ_BANK_KODI}},'),
    (1, 10, 2, 0, 'STIR: {{MIJOZ_STIR}},'),
    (1, 11, 2, 0, 'x/r: {{MIJOZ_HISOB_RAQAM}}.'),
    (1, 12, 2, 0, 'Rahbari: {{MIJOZ_RAHBAR}}'),
    (1, 13, 2, 0, 'Tel:{{MIJOZ_TEL}}'),
    (1, 14, 2, 0, '\xa0{{KAFIL_ISM}}'),
    (1, 15, 2, 0, '{{KAFIL_MANZIL}}'),
    (1, 16, 2, 0, 'PINFL: {{KAFIL_PINFL}}'),
    (1, 17, 2, 0, '{{KAFIL_PASSPORT_TOLIQ}}'),
    (1, 18, 2, 0, ''),
    (1, 19, 2, 0, 'Tel:{{KAFIL_TEL}}'),
]


def build_yuridik_kafil(input_path, output_path):
    doc = Document(input_path)
    for p in doc.paragraphs:
        _replace_in_paragraph_regex(p, PATTERNS_YURIDIK_KAFIL)
    for t_idx, r_idx, c_idx, p_idx, text in YURIDIK_KAFIL_TABLE_FIX:
        set_table_cell_paragraph(doc, t_idx, r_idx, c_idx, p_idx, text)
    doc.save(output_path)
    return output_path


PATTERNS_YURIDIK_KAFIL_GAROV = [
    (r'____________ 2026-yil.*?-\s*sonli', '{{ARIZA_SANA_QATORI}}'),
    (r'^Гулистон туманлараро иқтисодий судига\s*$', '{{SUD_NOMI}}'),
    (r'^Ўзбекистон Республикаси Савдо-саноат палатаси Сирдарё вилояти ҳудудий бошқармаси\s*$', '{{PALATA_NOMI}}'),
    (r'Даъвогар: \t“Агробанк” акциядорлик тижорат банки манфаатида ҳ\|р:\d+, банк коди:\d+, СТИР: \d+, Манзил:.*?10-уй\.',
     'Даъвогар: \t{{BANK_NOMI}} манфаатида ҳ|р:{{BANK_HISOB_RAQAM_BOSH}}, банк коди:{{BANK_KODI_BOSH}}, '
     'СТИР: {{BANK_STIR}}, Манзил: {{BANK_MANZIL_FILIAL}}.'),
    (r'Жавобгар:\t“SHODIXON OTA 2020” фермер ҳўжалиги х\\р:\d+, МФО:\d+,\s*ИНН:\d+, Юридик манзили:.*?СИУ\s*',
     'Жавобгар:\t{{MIJOZ_ISM}} х\\р:{{MIJOZ_HISOB_RAQAM}}, МФО:{{MIJOZ_BANK_KODI}}, ИНН:{{MIJOZ_STIR}}, '
     'Юридик манзили: {{MIJOZ_MANZIL}}'),
    (r'Қушимча жавобгар:\s+Каримов Абдурахмон Шодихонович\s+05\.01\.1982\s*',
     'Қушимча жавобгар:  {{KAFIL_ISM}}  {{KAFIL_TUGILGAN_SANA}} '),
    (r'Йил Боёвут туманида туғилган, ЖШШИР:\s*', 'йил {{KAFIL_TUGILGAN_JOY}}да туғилган, ЖШШИР: '),
    (r'30501822870042, манзили: Сирдарё вилояти,\s*', '{{KAFIL_PINFL}}, манзили: '),
    (r'Боёвут тумани, Миришкор МФЙ, Нуробод\s*', '{{KAFIL_MANZIL}} '),
    (r'кўчаси тел:\+99893-818-92-82', 'тел:{{KAFIL_TEL}}'),
    (r'“Агробанк” АТБ Боёвут филиали билан қарз олувчи “SHODIXON OTA 2020” фермер ҳўжалиги ўртасида тузилган .*? йилдаги №.*?-сонли кредит шартномасига биноан банк томонидан қарз олувчига .*? ой муддатга, .*? ой имтиёзли давр асосида, йиллик .*? фоиз устама ҳақи тўлаш шарти билан .*? мақсадида .*? сўм кредит маблағлари ажратиш келишилган ва шунча миқдордаги кредит маблағлари молиялаштирилган\.',
     '"{{BANK_QISQA_NOMI}}" АТБ {{FILIAL_NOMI}} филиали билан қарз олувчи {{MIJOZ_ISM}} ўртасида тузилган '
     '{{SHARTNOMA_SANA}} йилдаги кредит шартномасига биноан банк томонидан қарз олувчига '
     '{{KREDIT_MUDDATI_OY}} ой муддатга, {{IMTIYOZLI_DAVR_OY}} ой имтиёзли давр асосида, йиллик '
     '{{YILLIK_FOIZ}} фоиз устама ҳақи тўлаш шарти билан "{{KREDIT_MAQSAD}}" мақсадида {{KREDIT_SUMMA}} '
     'сўм кредит маблағлари ажратиш келишилган ва шунча миқдордаги кредит маблағлари молиялаштирилган.'),
    (r'Кредит таъминоти сифатида жавобгар Каримов Абдурахмон Шодихоновичнинг кафиллиги, шунингдек, жавобгар “SHODIXON OTA 2020” фермер ҳўжалигига тегишли .*? гаровга қўйилган\.\s*',
     'Кредит таъминоти сифатида жавобгар {{KAFIL_ISM}}нинг кафиллиги, шунингдек, {{GAROV_TAVSIFI}} '
     '{{GAROV_BAHOSI}} сўмга келишилиб гаровга қўйилган. '),
    (r'2026 йил 16 июль ҳолатига қарз олувчининг банк олдида .*? сўм муддати ўтган асосий кредит, .*? сўм кредитга хисобланган фоиз, .*? сўм кечиктирилган кунлар учун хисобланган пеня қарздорлиги вужудга келган, шунингдек, .*? сўм муддати келмаган асосий кредит қарздорлиги мавжуд, \.',
     '{{HOLAT_SANASI}} ҳолатига қарз олувчининг банк олдида {{ASOSIY_QARZ}} сўм муддати ўтган асосий '
     'кредит, {{FOIZ_QARZ}} сўм кредитга хисобланган фоиз, {{JARIMA}} сўм кечиктирилган кунлар учун '
     'хисобланган пеня қарздорлиги вужудга келган, шунингдек, {{MUDDATI_KELMAGAN_ASOSIY}} сўм '
     'муддати келмаган асосий кредит қарздорлиги мавжуд.'),
    (r'Жавобгар “SHODIXON OTA 2020” фермер ҳўжалиги ва кафил Каримов Абдурахмон Шодихонович хисобидан “Агробанк” АТБ Боёвут филиали фойдасига .*? сўм муддати ўтган асосий кредит, .*? сўм муддати келмаган асосий кредит, .*? сўм кредитга хисобланган фоиз, .*? сўм кечиктирилган кунлар учун хисобланган пеня қарздорлиги ва олдиндан тўланган почта ҳаражатини ундиришингизни;',
     'Жавобгар {{MIJOZ_ISM}} ва кафил {{KAFIL_ISM}} хисобидан "{{BANK_QISQA_NOMI}}" АТБ {{FILIAL_NOMI}} '
     'филиали фойдасига {{ASOSIY_QARZ}} сўм муддати ўтган асосий кредит, {{MUDDATI_KELMAGAN_ASOSIY}} сўм '
     'муддати келмаган асосий кредит, {{FOIZ_QARZ}} сўм кредитга хисобланган фоиз, {{JARIMA}} сўм '
     'кечиктирилган кунлар учун хисобланган пеня қарздорлиги ва олдиндан тўланган {{POCHTA_XARAJATI}} '
     'сўм почта ҳаражатини ундиришингизни;'),
    (r'Ундирувни “SHODIXON OTA 2020” фермер ҳўжалигига тегишли .*? гаровга қўйилган мол-мулкка қаратишингизни ва ушбу гаров мулкининг ким ошди савдосидаги бошланғич баҳосини .*? сўм деб белгилашингизни;',
     'Ундирувни {{GAROV_TAVSIFI}} гаровга қўйилган мол-мулкка қаратишингизни ва ушбу гаров мулкининг '
     'ким ошди савдосидаги бошланғич баҳосини {{GAROV_BAHOSI}} сўм деб белгилашингизни;'),
    (r'С\.Р\.Аманлиқов', '{{IMZO_ISM}}'),
    (r'Бошқарма бошлиғи\s*\n?\s*ўринбосари\s*:', '{{IMZO_LAVOZIM}}:'),
]


# =====================================================================
# 6) Muddatidan oldin undirish (shartnomani bekor qilish) — juda holatga
#    xos hujjat, faqat umumiy qismlar parametrlanadi, tafsilotlar
#    "qo'shimcha izoh" maydonidan qo'lda kiritiladi.
# =====================================================================
PATTERNS_MUDDATIDAN_OLDIN = [
    (r'____________ 2026-yil.*?-\s*sonli', '{{ARIZA_SANA_QATORI}}'),
    (r'^Гулистон туманлараро иқтисодий судига\s*$', '{{SUD_NOMI}}'),
    (r'^Ўзбекистон Республикаси Савдо-саноат палатаси Сирдарё вилояти ҳудудий бошқармаси\s*$', '{{PALATA_NOMI}}'),
    (r'Даъвогар: \t“Агробанк” акциядорлик тижорат банки манфаатида ҳ\|р:\d+, банк коди:\d+, СТИР: \d+, Манзил:.*?10-уй\.',
     'Даъвогар: \t{{BANK_NOMI}} манфаатида ҳ|р:{{BANK_HISOB_RAQAM_BOSH}}, банк коди:{{BANK_KODI_BOSH}}, '
     'СТИР: {{BANK_STIR}}, Манзил: {{BANK_MANZIL_FILIAL}}.'),
    (r'Жавобгар:\t“AGRO SABZAVOT” фермер ҳўжалиги х\\р: \d+, МФО: \d+,\s*ИНН: \d+, Юридик манзили:.*?худуди\s*',
     'Жавобгар:\t{{MIJOZ_ISM}} х\\р: {{MIJOZ_HISOB_RAQAM}}, МФО: {{MIJOZ_BANK_KODI}}, ИНН: {{MIJOZ_STIR}}, '
     'Юридик манзили: {{MIJOZ_MANZIL}}'),
    (r'“Агробанк” акциядорлик тижорат банки \(Сирдарё вилоят Боёвут филиали\) билан қарз олувчи “AGRO SABZAVOT” фермер ҳўжалиги ўртасида тузилган .*?йилдаги №.*?-сонли кредит шартномасига биноан.*?ажратиши белгиланган\.\s*',
     '{{VOQEALAR_TAVSIFI}}\n\n'),
    (r'Буюртмачи “AGRO SABZAVOT”.*?якунланишини таъминлаш белгиланган\.\s*', ''),
    (r'Қарз олувчи томонидан тақдим қилинган тўлов топширқномасига асосан.*?ўтказиб берилган\.\s*', ''),
    (r'Банк томонидан кредит маблағлари мақсадли сарфланиши ҳолати мониторинг қилинганда.*?тасдиқланган\.\s*', ''),
    (r'Мазкур ҳолатда қарз олувчи томонидан кредитдан фойдаланиш даврида.*?йўл қўйган\.',
     'Мазкур ҳолатда қарз олувчи томонидан кредитдан фойдаланиш даврида ўз зиммасига олган '
     'мажбуриятларни бажармаган.'),
    (r'Қарз олувчи томонидан кредит маблағлари мақсадсиз сарфланганлиги ва шартнома шартлари жиддий равишда бузилганлиги сабабли банк томонидан кредит шартномасини муддатидан олдин 00 кун муддат ичида бекор қилиш юзасидан 00\.00\.2026 йилда жавобгарга таклиф хати берилган бўлсада, жавобгар томонидан ушбу таклиф хати эътиборсиз қолдирилган\.',
     'Қарз олувчи томонидан шартнома шартлари жиддий равишда бузилганлиги сабабли банк томонидан '
     'кредит шартномасини муддатидан олдин бекор қилиш юзасидан {{XAT_SANASI}}да жавобгарга '
     '{{XAT_TURI_NOMI}} берилган бўлсада, жавобгар томонидан ушбу хат эътиборсиз қолдирилган.'),
    (r'2026 йил 26 февраль ҳолатига қарз олувчининг банк олдида .*? сўм муддати келмаган асосий кредит, .*? сўм кредитга хисобланган фоиз, .*? сўм кечиктирилган кунлар учун хисобланган пеня қарздорлиги мавжуд\.',
     '{{HOLAT_SANASI}} ҳолатига қарз олувчининг банк олдида {{MUDDATI_KELMAGAN_ASOSIY}} сўм муддати '
     'келмаган асосий кредит, {{FOIZ_QARZ}} сўм кредитга хисобланган фоиз, {{JARIMA}} сўм кечиктирилган '
     'кунлар учун хисобланган пеня қарздорлиги мавжуд.'),
    (r'Кредит шартномасининг 10\.3-бандида,.*?ошиб кетмаслиги назарда тутилган\.', ''),
    (r'“Агробанк” акциядорлик тижорат банки билан жавобгар “AGRO SABZAVOT” фермер ҳўжалиги ўртасида тузилган .*?йилдаги №.*?-сонли кредит шартномасини муддатидан олдин бекор қилишни;',
     '{{BANK_QISQA_NOMI}} акциядорлик тижорат банки билан жавобгар {{MIJOZ_ISM}} ўртасида тузилган '
     '{{SHARTNOMA_SANA}}йилдаги кредит шартномасини муддатидан олдин бекор қилишни;'),
    (r'“AGRO SABZAVOT” фермер ҳўжалиги хисобидан “Агробанк” акциядорлик тижорат банки фойдасига .*? сўм муддати келмаган асосий кредит, .*? сўм кредитга хисобланган фоиз, .*? сўм кечиктирилган кунлар учун хисобланган пеня қарздорлигини ундиришингизни;',
     '{{MIJOZ_ISM}} хисобидан {{BANK_QISQA_NOMI}} акциядорлик тижорат банки фойдасига '
     '{{MUDDATI_KELMAGAN_ASOSIY}} сўм муддати келмаган асосий кредит, {{FOIZ_QARZ}} сўм кредитга '
     'хисобланган фоиз, {{JARIMA}} сўм кечиктирилган кунлар учун хисобланган пеня қарздорлигини '
     'ундиришингизни;'),
    (r'Сирдарё вилояти, Боёвут тумани, Боёвут шаҳарчаси, Тинчлик кўчаси, 10-уйда жойлашган банкнинг Сирдарё вилоят бошқармаси Боёвут филиалига',
     '{{BANK_MANZIL_FILIAL}} жойлашган {{FILIAL_NOMI}} филиалига'),
    (r'А\.Ш\.Мансуров', '{{IMZO_ISM}}'),
    (r'Бошқарма бошлиғи:', '{{IMZO_LAVOZIM}}:'),
]

if __name__ == '__main__':
    import os
    OUT = '/home/claude/qarz_nazorat/templates'
    os.makedirs(OUT, exist_ok=True)

    templatize('Jismoniy_shaxslarga_davo_ariza.docx',
               f'{OUT}/davo_jismoniy_oddiy.docx', PATTERNS_JISMONIY_ODDIY)
    templatize_by_index('Jismoniy_shaxslarning_kafilikka_qaratish.docx',
                         f'{OUT}/davo_jismoniy_kafil.docx', JISMONIY_KAFIL_INDEX_FIX,
                         patterns=PATTERNS_JISMONIY_KAFIL)
    templatize('Iqtisodiy_sudga.docx',
               f'{OUT}/davo_yuridik_oddiy.docx', PATTERNS_YURIDIK_ODDIY)
    # СТИР maydoni band+mijoz uchun bir xil ko'rinishda edi — mijoznikini alohida to'g'irlaymiz
    _fix_doc = Document(f'{OUT}/davo_yuridik_oddiy.docx')
    set_table_cell_paragraph(_fix_doc, 1, 2, 2, 3, 'СТИР: {{MIJOZ_STIR}},')
    _fix_doc.save(f'{OUT}/davo_yuridik_oddiy.docx')
    build_yuridik_kafil('Kafillikka_qaratish_va_foiz_penya_undirsh_haqida.docx',
                         f'{OUT}/davo_yuridik_kafil.docx')
    templatize('Kafillik_va_garovga_qaratish_togrisida.docx',
               f'{OUT}/davo_yuridik_kafil_garov.docx', PATTERNS_YURIDIK_KAFIL_GAROV)
    templatize('Muddatidan_oldin_undirish_uchun.docx',
               f'{OUT}/davo_muddatidan_oldin.docx', PATTERNS_MUDDATIDAN_OLDIN)
    print("6 ta shablon yaratildi.")
    print("6 ta shablon yaratildi.")
