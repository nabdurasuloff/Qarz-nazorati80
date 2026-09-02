# -*- coding: utf-8 -*-
"""
Yuklangan haqiqiy Davo ariza namunalarini shablonga aylantirish uchun
umumiy yordamchi: har bir paragraf/jadval katagi matnini butunlay o'qib,
aniq (exact) matn parchalarini {{PLACEHOLDER}} bilan almashtiradi, so'ng
formatlashni saqlagan holda birinchi run'ga yozadi.
"""
from docx import Document


def _replace_in_paragraph(paragraph, replacements):
    full_text = ''.join(run.text for run in paragraph.runs)
    if not full_text:
        return
    new_text = full_text
    for old, new in replacements:
        if old in new_text:
            new_text = new_text.replace(old, new)
    if new_text == full_text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def templatize(input_path, output_path, replacements):
    """
    replacements: list of (old_exact_text, new_placeholder_text) tuples.
    Uzunroq matnlar oldin almashtirilishi kerak (qisqa matnlar ularning
    ichida bo'lib qolmasligi uchun) — chaqiruvchi tomondan tartiblangan
    bo'lishi kerak.
    """
    doc = Document(input_path)
    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, replacements)
    doc.save(output_path)
    return output_path


def dump_text(path):
    doc = Document(path)
    lines = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            lines.append(f"P{i}: {p.text}")
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if cell.text.strip():
                    lines.append(f"T{ti}R{ri}C{ci}: {cell.text}")
    return '\n'.join(lines)
