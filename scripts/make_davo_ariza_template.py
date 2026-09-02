# -*- coding: utf-8 -*-
"""
Davo ariza (sudga qarzni undirish to'g'risida) uchun Word shablonini yaratadi.
DIQQAT: bu standart tuzilma asosida tayyorlangan LOYIHA shablon — yuridik
bo'limingiz tomonidan tekshirilib, kerak bo'lsa moslashtirilishi tavsiya etiladi.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_border_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para(doc, text='', size=11, bold=False, italic=False, align=None,
             space_after=8, font_name='Times New Roman', color=None,
             indent_first=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if indent_first is not None:
        p.paragraph_format.first_line_indent = Cm(indent_first)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # ---- Sud nomi va taraflar (o'ng tomonda, "shapka" uslubida) ----
    add_para(doc, "{{SUD_NOMI}}", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)

    add_para(doc, "Даъвогар:", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "{{BANK_NOMI}}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "{{FILIAL_NOMI}} филиали", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "Манзил: {{BANK_MANZIL}}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "Тел: {{FILIAL_TEL}}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)

    add_para(doc, "Жавобгар:", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "{{MIJOZ_ISM}}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "Манзил: {{MIJOZ_MANZIL}}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_para(doc, "Ҳужжат/СТИР: {{MIJOZ_HUJJAT}}", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14)

    add_para(doc, "Даъво суммаси: {{JAMI_QARZ}} сўм", size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    add_para(doc, "Давлат божи: {{DAVLAT_BOJI}} сўм", size=12,
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=24)

    # ---- Sarlavha ----
    add_para(doc, "ДАЪВО АРИЗАСИ", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "(кредит шартномаси бўйича қарздорликни ундириш тўғрисида)",
             size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # ---- Asosiy matn ----
    add_para(
        doc,
        '"{{BANK_QISQA_NOMI}}" АТБ {{FILIAL_NOMI}} филиали (бундан кейин — "Даъвогар" деб '
        'юритилади) билан {{MIJOZ_ISM}} (бундан кейин — "Жавобгар" деб юритилади) ўртасида '
        '{{SHARTNOMA_SANA}}да тузилган {{ANKETA_RAQAM}}-сонли кредит шартномасига асосан '
        'Жавобгарга {{KREDIT_SUMMA}} сўм миқдорида {{KREDIT_MAQSAD}} кредити ажратилган.',
        size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=12
    )

    add_para(
        doc,
        'Шартнома шартларига мувофиқ Жавобгар кредит маблағини ва унга ҳисобланган '
        'фоизларни белгиланган жадвал асосида қайтариш мажбуриятини олган. Бироқ '
        'Жавобгар томонидан ушбу мажбурият лозим даражада бажарилмади, натижада '
        '{{HOLAT_SANASI}} ҳолатига жами муддати ўтган қарздорлик {{JAMI_QARZ}} сўмни '
        'ташкил этмоқда, жумладан: муддати ўтган асосий қарз — {{ASOSIY_QARZ}} сўм, '
        'муддати ўтган фоиз — {{FOIZ_QARZ}} сўм, ҳисобланган жарима (пеня) — '
        '{{JARIMA}} сўм.',
        size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=12
    )

    add_para(
        doc,
        'Даъвогар томонидан Жавобгарга {{XAT_SANASI}}да {{XAT_TURI_NOMI}} йўлланган бўлиб, '
        'унда қарздорликни кўрсатилган муддат ичида тўлаш талаб қилинган. Бироқ Жавобгар '
        'томонидан белгиланган муддатда қарздорлик тўланмади ва/ёки жавоб берилмади.',
        size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=12
    )

    add_para(
        doc,
        'Юқоридагиларга асосан ва Ўзбекистон Республикаси Фуқаролик кодексининг '
        '(қарз мажбуриятларини бажариш, кредит шартномаси) тегишли моддаларига мувофиқ, '
        'Даъвогар қуйидагиларни сўрайди:',
        size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=1.25, space_after=14
    )

    for i, text in enumerate([
        'Жавобгар {{MIJOZ_ISM}}дан {{ANKETA_RAQAM}}-сонли кредит шартномаси бўйича '
        'муддати ўтган {{JAMI_QARZ}} сўм қарздорликни (шу жумладан асосий қарз '
        '{{ASOSIY_QARZ}} сўм, фоиз {{FOIZ_QARZ}} сўм, жарима {{JARIMA}} сўм) Даъвогар '
        'фойдасига ундирилсин;',
        'Ушбу иш бўйича тўланган давлат божи ({{DAVLAT_BOJI}} сўм) ва бошқа суд '
        'харажатлари Жавобгар зиммасига юклатилсин.',
    ], start=1):
        add_para(doc, f'{i}. {text}', size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  indent_first=1.25, space_after=10)

    add_para(doc, 'Иловалар:', size=12, bold=True, space_after=6)
    for item in [
        'Кредит шартномаси нусхаси;',
        '{{XAT_TURI_NOMI}} нусхаси ва уни юборилганлигини тасдиқловчи ҳужжат;',
        'Қарздорлик ҳисоб-китоби (асосий қарз, фоиз, жарима бўйича);',
        'Давлат божи тўланганлигини тасдиқловчи ҳужжат;',
        'Ушбу аризанинг Жавобгарга юборилганлигини тасдиқловчи ҳужжат.',
    ]:
        add_para(doc, f'—  {item}', size=11.5, space_after=4)

    add_para(doc, '', size=6, space_after=20)

    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.columns[0].width = Cm(10)
    sign_table.columns[1].width = Cm(6.5)
    lcell, rcell = sign_table.rows[0].cells
    set_cell_border_none(lcell)
    set_cell_border_none(rcell)

    lp = lcell.paragraphs[0]
    r1 = lp.add_run('Сана: {{ARIZA_SANASI}}')
    r1.font.size = Pt(12)

    rp = rcell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = rp.add_run('"{{BANK_QISQA_NOMI}}" АТБ\n{{FILIAL_NOMI}} филиали бошқарувчиси\n\n{{RAHBAR_ISM}}')
    r2.bold = True
    r2.font.size = Pt(12)

    doc.save('/home/claude/qarz_nazorat/templates/davo_ariza_shablon.docx')
    print("Davo ariza shabloni yaratildi.")


if __name__ == '__main__':
    build()
