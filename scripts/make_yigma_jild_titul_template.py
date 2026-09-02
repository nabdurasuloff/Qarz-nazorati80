# -*- coding: utf-8 -*-
"""
Ijro harakati (MIB) yig'ma jildi uchun standart titul (muqova) shablonini
yaratadi. Bu — standart shablon; foydalanuvchi keyinchalik o'zining
rasmiy shablonini yuklab, bu faylni Sozlamalar orqali almashtirishi mumkin.
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_centered(doc, text, size=11, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


doc = Document()

add_centered(doc, "{{BANK_NOMI}}", size=12, bold=True)
add_centered(doc, "{{FILIAL_NOMI}} filiali", size=11)

doc.add_paragraph()
doc.add_paragraph()

add_centered(doc, "IJRO HARAKATI\nYIG'MA JILDI", size=20, bold=True)

doc.add_paragraph()

add_centered(doc, "Ish raqami: {{MIB_ISH_RAQAMI}}", size=13, bold=True)

doc.add_paragraph()
doc.add_paragraph()

table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True


def qator(label, val):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = val
    for cell in row.cells:
        for para in cell.paragraphs:
            for r in para.runs:
                r.font.size = Pt(11)
    if row.cells[0].paragraphs[0].runs:
        row.cells[0].paragraphs[0].runs[0].bold = True


qator("Qarzdor", "{{QARZDOR_ISM}}")
qator("Turi", "{{MIJOZ_TURI}}")
qator("Manzili", "{{MIJOZ_MANZIL}}")
qator("Hujjat raqami (PINFL/STIR)", "{{HUJJAT_RAQAMI}}")
qator("Anketa raqami", "{{ANKETA_RAQAM}}")
qator("Jami qarzdorlik (so'm)", "{{JAMI_QARZ}}")
qator("", "")
qator("Ogohlantirish/Talabnoma xati sanasi", "{{XAT_SANASI}}")
qator("Xat yuborilgan sana", "{{XAT_YUBORILGAN_SANA}}")
qator("Da'vo ariza tayyorlangan sana", "{{DAVO_ARIZA_SANA}}")
qator("Palatadan qaytgan sana", "{{PALATADAN_QAYTGAN_SANA}}")
qator("Sudga topshirilgan sana", "{{SUDGA_TOPSHIRILGAN_SANA}}")
qator("Sud ishi raqami", "{{SUD_ISH_RAQAMI}}")
qator("", "")
qator("MIBga o'tkazilgan sana", "{{MIB_OTKAZILGAN_SANA}}")
qator("MIB ish raqami", "{{MIB_ISH_RAQAMI}}")
qator("Jild ochilgan sana", "{{JILD_OCHILGAN_SANA}}")

doc.add_paragraph()
doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
frun = footer.add_run("Ushbu jildga barcha tegishli hujjatlar (ijro varaqasi, dalolatnoma, "
                       "sotuv hujjatlari va boshqalar) tartib bilan ilova qilib boriladi.")
frun.italic = True
frun.font.size = Pt(9)

doc.save('/home/claude/qarz_nazorat/templates/yigma_jild_titul_shablon.docx')
print("Shablon yaratildi")
