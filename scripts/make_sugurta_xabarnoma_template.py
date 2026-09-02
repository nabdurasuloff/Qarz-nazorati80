# -*- coding: utf-8 -*-
"""
Sug'urta kompaniyasiga yuboriladigan xabarnoma uchun standart Word shablonini
yaratadi. Bu — vaqtinchalik/standart shablon; foydalanuvchi keyinchalik
o'zining haqiqiy shablonini yuklab, bu faylni almashtirishi mumkin
(Sozlamalar bo'limidan).
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_para(doc, text='', size=11, bold=False, align=None, space_after=8):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p


doc = Document()

add_para(doc, "{{BANK_NOMI}}", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "{{FILIAL_NOMI}} филиали", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "")

add_para(doc, "{{SUGURTA_KOMPANIYA}}га", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para(doc, "")

add_para(doc, "№ {{XABARNOMA_SANA}}", size=11)
add_para(doc, "")

add_para(doc, "ХАБАРНОМА", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "(қарз олувчининг вафот этиши муносабати билан суғурта "
              "тўлови тўғрисида)", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "")

add_para(
    doc,
    "\"{{BANK_QISQA_NOMI}}\" АТБ {{FILIAL_NOMI}} филиали билан {{MIJOZ_ISM}} "
    "(ПИНФЛ: {{MIJOZ_PINFL}}) ўртасида {{SHARTNOMA_SANA}} йилда тузилган "
    "кредит шартномасига (анкета рақами: {{ANKETA_RAQAM}}) биноан "
    "{{KREDIT_SUMMA}} сўм миқдорида кредит маблағи ажратилган эди.",
    space_after=10
)

add_para(
    doc,
    "Афсуски, қарз олувчи {{MIJOZ_ISM}} {{VAFOT_SANASI}} санада вафот этгани "
    "тўғрисида маълумот олинди (васиқа/гувоҳнома нусхаси илова қилинмоқда).",
    space_after=10
)

add_para(
    doc,
    "Юқоридаги кредит шартномаси бўйича {{SUGURTA_KOMPANIYA}} томонидан "
    "{{SUGURTA_POLIS_RAQAM}} рақамли суғурта полиси расмийлаштирилган бўлиб, "
    "ушбу полис амал қилиш муддати {{KREDIT_TUGASH_SANASI}} санагача давом "
    "этади.",
    space_after=10
)

add_para(
    doc,
    "Юқоридагиларни инобатга олиб, қарз олувчининг вафоти муносабати билан "
    "юзага келган {{JAMI_QARZ}} сўм миқдоридаги кредит қарздорлигини "
    "суғурта шартномаси асосида қоплаб беришингизни сўраймиз.",
    space_after=10
)

add_para(doc, "Илова: ўлим тўғрисидаги гувоҳнома нусхаси, паспорт нусхаси, "
              "кредит шартномаси нусхаси.", size=10)
add_para(doc, "")
add_para(doc, "")

add_para(doc, "Филиал бошқарувчиси:" + " " * 30 + "{{RAHBAR_ISM}}", size=11)

doc.save('/home/claude/qarz_nazorat/templates/sugurta_xabarnoma_shablon.docx')
print("Shablon yaratildi")
