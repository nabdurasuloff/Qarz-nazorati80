# -*- coding: utf-8 -*-
"""
Word shabloniga ma'lumotlarni joylab, tayyor xat (.docx) yaratish.
"""
import os
import re
import sys
import datetime
from docx import Document
import util


def _base_dir():
    # PyInstaller --onefile bilan yig'ilganda fayllar vaqtinchalik papkaga
    # (sys._MEIPASS) ochiladi; oddiy ishga tushirishda esa shu faylning papkasi.
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        return sys._MEIPASS
    return os.path.dirname(os.path.abspath(__file__))


TEMPLATE_PATH = os.path.join(_base_dir(), 'templates', 'xat_shablon.docx')
DAVO_ARIZA_TEMPLATE_PATH = os.path.join(_base_dir(), 'templates', 'davo_ariza_shablon.docx')
SUGURTA_XABARNOMA_TEMPLATE_PATH = os.path.join(_base_dir(), 'templates', 'sugurta_xabarnoma_shablon.docx')
YIGMA_JILD_TITUL_TEMPLATE_PATH = os.path.join(_base_dir(), 'templates', 'yigma_jild_titul_shablon.docx')


def _fmt_summa(val):
    try:
        val = float(val)
    except (TypeError, ValueError):
        return str(val)
    return f"{val:,.0f}".replace(',', ' ')


def _replace_in_paragraph(paragraph, mapping):
    full_text = ''.join(run.text for run in paragraph.runs)
    if '{{' not in full_text:
        return
    new_text = full_text
    for key, val in mapping.items():
        # Katta-kichik harf farqiga sezgir bo'lmagan almashtirish — agar
        # shablon biror joyda (masalan Word orqali) qo'lda tahrirlanib,
        # placeholder harflari o'zgarib qolgan bo'lsa ham ishlashi uchun.
        pattern = re.compile(r'\{\{\s*' + re.escape(key) + r'\s*\}\}', re.IGNORECASE)
        new_text = pattern.sub(str(val), new_text)
    if new_text == full_text:
        return
    # Barcha runlarni tozalab, birinchi runga yangi matnni yozamiz
    # (formatlashni birinchi run’dan olamiz)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def _replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def generate_letter(output_path, xat_turi, mijoz_ism, mijoz_manzil, portfel_row, settings,
                     anketa_raqami=None, rahbar_ism=None):
    """
    xat_turi: 'Ogohlantirish' yoki 'Talabnoma'
    portfel_row: dict — bitta portfel qatoridagi ma'lumot (database.py formatida)
    settings: dict — get_all_settings() natijasi
    """
    doc = Document(TEMPLATE_PATH)

    sarlavha = "OGOHLANTIRISH XATI" if xat_turi == 'Ogohlantirish' else "TALABNOMA"

    holat_sanasi = datetime.date.today().strftime('%d.%m.%Y')

    mapping = {
        'BANK_NOMI': settings.get('bank_nomi', ''),
        'BANK_QISQA_NOMI': settings.get('bank_qisqa_nomi', ''),
        'BANK_MANZIL': settings.get('bank_manzil', ''),
        'BANK_EMAIL': settings.get('bank_email', ''),
        'BANK_SAYT': settings.get('bank_sayt', ''),
        'BANK_TEL': settings.get('bank_tel', ''),
        'BANK_MOBIL_ILOVA': settings.get('bank_mobil_ilova', ''),
        'BANK_KODI': settings.get('bank_kodi', ''),
        'ALOQA_MARKAZI_TEL': settings.get('aloqa_markazi_tel', ''),
        'FILIAL_NOMI': settings.get('filial_nomi', ''),
        'FILIAL_TEL': settings.get('filial_tel', ''),
        'RAHBAR_ISM': rahbar_ism or settings.get('rahbar_ism', ''),

        'MIJOZ_ISM': mijoz_ism or '',
        'MIJOZ_MANZIL': mijoz_manzil or '',
        'SARLAVHA': sarlavha,

        'KREDIT_SUMMA': _fmt_summa(portfel_row.get('jami_berilgan_summa') or portfel_row.get('ead', 0)),
        'KREDIT_MAQSAD': portfel_row.get('tulov_maqsadi', '') or '',
        'HOLAT_SANASI': holat_sanasi,
        'JAMI_QARZ': _fmt_summa(portfel_row.get('jami_qarz', 0)),
        'ASOSIY_QARZ': _fmt_summa(portfel_row.get('asosiy_qarz', 0)),
        'FOIZ_QARZ': _fmt_summa(portfel_row.get('foiz_qarz', 0)),
        'JARIMA': _fmt_summa(portfel_row.get('jarima', 0)),
        'TOLOV_MUDDATI': f"{settings.get('tolov_muddati_kun', '10')} bank ish kuni",
        'ANKETA_RAQAM': anketa_raqami or portfel_row.get('anketa_raqami', ''),
    }

    _replace_everywhere(doc, mapping)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_davo_ariza(*args, **kwargs):
    raise NotImplementedError(
        "Bu funksiya eskirgan. O'rniga generate_davo_ariza_v2() dan foydalaning."
    )


def safe_filename(text):
    text = re.sub(r'[\\/*?:"<>|]', '', str(text))
    text = text.strip().replace(' ', '_')
    return text[:80]


DAVO_ARIZA_TEMPLATES = {
    'jismoniy_oddiy': 'davo_jismoniy_oddiy.docx',
    'jismoniy_kafil': 'davo_jismoniy_kafil.docx',
    'jismoniy_garov': 'davo_jismoniy_garov.docx',
    'jismoniy_kafil_garov': 'davo_jismoniy_kafil_garov.docx',
    'yuridik_oddiy': 'davo_yuridik_oddiy.docx',
    'yuridik_kafil': 'davo_yuridik_kafil.docx',
    'yuridik_kafil_garov': 'davo_yuridik_kafil_garov.docx',
    'yuridik_garov': 'davo_yuridik_garov.docx',
    'muddatidan_oldin': 'davo_muddatidan_oldin.docx',
}

DAVO_ARIZA_NOMLARI = {
    'jismoniy_oddiy': "Jismoniy shaxs — oddiy (kafilsiz)",
    'jismoniy_kafil': "Jismoniy shaxs — kafil bilan",
    'jismoniy_garov': "Jismoniy shaxs — faqat garov mulkiga qaratish (kafilsiz)",
    'jismoniy_kafil_garov': "Jismoniy shaxs — kafil + garov mulkiga qaratish",
    'yuridik_oddiy': "Yuridik/F.X — oddiy (kafilsiz)",
    'yuridik_kafil': "Yuridik shaxs — kafil bilan (foiz-penya)",
    'yuridik_kafil_garov': "Yuridik shaxs — kafil + garov mulkiga qaratish",
    'yuridik_garov': "Yuridik shaxs — faqat garov mulkiga qaratish (kafilsiz)",
    'muddatidan_oldin': "Shartnomani bekor qilish (muddatidan oldin undirish)",
}


def tavsiya_ariza_turi(mijoz_turi, taminot):
    """
    Mijoz turi va kiritilgan ta'minot (kafillik/garov) ma'lumotiga qarab,
    eng mos Davo ariza turini tavsiya qiladi.
    taminot: davo_taminot jadvalidagi dict (yoki None/bo'sh).
    """
    taminot_turi = (taminot or {}).get('taminot_turi') or 'yoq'
    # YaTT (yakka tartibdagi tadbirkor) — shaxsan jismoniy shaxs bo'lsa-da,
    # tadbirkorlik faoliyati yuritgani uchun Davo ariza turi tanlashda
    # yuridik shaxs kabi ko'riladi (iqtisodiy sud, tegishli shablon).
    if mijoz_turi in ('yuridik', 'yatt'):
        if taminot_turi == 'kafillik_garov':
            return 'yuridik_kafil_garov'
        elif taminot_turi == 'kafillik':
            return 'yuridik_kafil'
        elif taminot_turi == 'garov':
            return 'yuridik_garov'
        return 'yuridik_oddiy'
    else:
        if taminot_turi == 'kafillik_garov':
            return 'jismoniy_kafil_garov'
        elif taminot_turi == 'kafillik':
            return 'jismoniy_kafil'
        elif taminot_turi == 'garov':
            return 'jismoniy_garov'
        return 'jismoniy_oddiy'


def _oy_farqi(sana1_str, sana2_str):
    """Ikki sana orasidagi farqni oy hisobida qaytaradi (taxminiy)."""
    for fmt_pair in [('%d.%m.%Y', '%d.%m.%Y')]:
        try:
            d1 = datetime.datetime.strptime(str(sana1_str)[:10], '%d.%m.%Y')
            d2 = datetime.datetime.strptime(str(sana2_str)[:10], '%d.%m.%Y')
            months = (d2.year - d1.year) * 12 + (d2.month - d1.month)
            return max(months, 0)
        except (ValueError, TypeError):
            return ''
    return ''


def generate_davo_ariza_v2(turi, output_path, portfel_row, mijoz, taminot, settings,
                            xat_sanasi='', xat_turi_nomi='', imzo_ism_override=None):
    """
    turi: DAVO_ARIZA_TEMPLATES kalitlaridan biri.
    mijoz: mijozlar jadvalidagi dict (yoki None).
    taminot: davo_taminot jadvalidagi dict (yoki None) — kafil/garov ma'lumotlari.
    """
    tpl_file = DAVO_ARIZA_TEMPLATES.get(turi)
    if not tpl_file:
        raise ValueError(f"Noma'lum davo ariza turi: {turi}")
    tpl_path = os.path.join(_base_dir(), 'templates', tpl_file)
    doc = Document(tpl_path)

    mijoz = mijoz or {}
    taminot = taminot or {}
    holat_sanasi = datetime.date.today().strftime('%d.%m.%Y')
    muddat_oy = _oy_farqi(portfel_row.get('shartnoma_sanasi'), portfel_row.get('shartnoma_tugash_sanasi'))

    jami_ead = float(portfel_row.get('ead', 0) or 0)
    asosiy = float(portfel_row.get('asosiy_qarz', 0) or 0)
    foiz = float(portfel_row.get('foiz_qarz', 0) or 0)
    jarima_v = float(portfel_row.get('jarima', 0) or 0)
    jami = asosiy + foiz + jarima_v
    muddati_kelmagan = max(jami_ead - jami, 0)

    pochta = taminot.get('pochta_xarajati') or settings.get('pochta_xarajati_standart', '41200')

    mapping = {
        'ARIZA_SANA_QATORI': f'{holat_sanasi}-yil',
        'SUD_NOMI': settings.get('sud_iqtisodiy_nomi' if turi.startswith('yuridik') or turi == 'muddatidan_oldin'
                                  else 'sud_fuqarolik_nomi', ''),
        'PALATA_NOMI': settings.get('palata_nomi', ''),

        'BANK_NOMI': settings.get('bank_nomi', ''),
        'BANK_QISQA_NOMI': settings.get('bank_qisqa_nomi', ''),
        'BANK_STIR': settings.get('bank_stir', ''),
        'BANK_KODI_BOSH': settings.get('bank_kodi_bosh', ''),
        'BANK_HISOB_RAQAM_BOSH': settings.get('bank_hisob_raqami_bosh', ''),
        'BANK_MANZIL_BOSH': settings.get('bank_manzil', ''),
        'BANK_MANZIL_FILIAL': settings.get('bank_rasmiy_manzil_filial', ''),
        'BANK_REKVIZIT_QATORI': (f"ҳ|р: {settings.get('bank_hisob_raqami_filial', '')}, "
                                  f"банк коди: {settings.get('bank_kodi_filial', '')}, "
                                  f"СТИР: {settings.get('bank_stir', '')}, "
                                  f"Манзил: {settings.get('bank_rasmiy_manzil_filial', '')}"),
        'FILIAL_NOMI': settings.get('filial_nomi', ''),

        'MIJOZ_ISM': util.mijoz_ism_hujjat_uchun(
            mijoz.get('ism') or portfel_row.get('mijoz_nomi', ''),
            util.turi_kodidan(portfel_row.get('mijoz_turi_kodi'), portfel_row.get('mijoz_turi'))
        ),
        'MIJOZ_MANZIL': mijoz.get('manzil', '') or '',
        'MIJOZ_PINFL': portfel_row.get('pinfl', '') or mijoz.get('hujjat_raqami', '') or '',
        'MIJOZ_STIR': portfel_row.get('stir', '') or '',
        'MIJOZ_BANK_KODI': portfel_row.get('filial_kodi', '') or '',
        'MIJOZ_HISOB_RAQAM': portfel_row.get('kredit_hisob_raqami', '') or '',
        'MIJOZ_RAHBAR': mijoz.get('rahbar_ism', '') or '',
        'MIJOZ_TEL': mijoz.get('telefon', '') or '',
        'MIJOZ_PASSPORT_TOLIQ': _passport_toliq(mijoz),
        'BANK_HUDUDIY_NOMI': f"{settings.get('viloyat_nomi','Сирдарё')} вилоят худудий бошкармаси ({settings.get('filial_nomi','')} филиали)",

        'KAFIL_ISM': taminot.get('kafil_ism', '') or '',
        'KAFIL_MANZIL': taminot.get('kafil_manzil', '') or '',
        'KAFIL_PINFL': taminot.get('kafil_pinfl', '') or '',
        'KAFIL_TEL': taminot.get('kafil_tel', '') or '',
        'KAFIL_TUGILGAN_SANA': taminot.get('kafil_passport_sana', '') or '',
        'KAFIL_TUGILGAN_JOY': taminot.get('kafil_passport_organ', '') or '',
        'KAFIL_PASSPORT_TOLIQ': (
            f"{taminot.get('kafil_passport','')}, {taminot.get('kafil_passport_sana','')} йил "
            f"{taminot.get('kafil_passport_organ','')}дан берилган" if taminot.get('kafil_passport') else ''
        ),

        'GAROV_TAVSIFI': taminot.get('garov_tavsifi', '') or '',
        'GAROV_BAHOSI': _fmt_summa(taminot.get('garov_bahosi', 0)),

        'SHARTNOMA_SANA': portfel_row.get('shartnoma_sanasi', '') or '',
        'KREDIT_MUDDATI_OY': muddat_oy,
        'IMTIYOZLI_DAVR_OY': '',
        'YILLIK_FOIZ': portfel_row.get('yillik_foiz', '') or '',
        'KREDIT_MAQSAD': portfel_row.get('tulov_maqsadi', '') or '',
        'KREDIT_SUMMA': _fmt_summa(portfel_row.get('jami_berilgan_summa') or portfel_row.get('ead', 0)),

        'HOLAT_SANASI': holat_sanasi,
        'JAMI_QARZ': _fmt_summa(jami),
        'ASOSIY_QARZ': _fmt_summa(asosiy),
        'FOIZ_QARZ': _fmt_summa(foiz),
        'JARIMA': _fmt_summa(jarima_v),
        'MUDDATI_KELMAGAN_ASOSIY': _fmt_summa(muddati_kelmagan),
        'POCHTA_XARAJATI': _fmt_summa(pochta),

        'XAT_SANASI': xat_sanasi,
        'XAT_TURI_NOMI': xat_turi_nomi,
        'VOQEALAR_TAVSIFI': taminot.get('garov_tavsifi', '') or '[Voqealar tavsifini shu yerga kiriting]',

        'IMZO_ISM': imzo_ism_override or settings.get('sud_ariza_imzo_ism', '') or '',
        'IMZO_LAVOZIM': settings.get('sud_ariza_imzo_lavozimi', ''),
    }

    _replace_everywhere(doc, mapping)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _passport_toliq(mijoz):
    if not mijoz:
        return ''
    hujjat = mijoz.get('hujjat_raqami', '') or ''
    sana = mijoz.get('passport_sana', '') or ''
    organ = mijoz.get('passport_organ', '') or ''
    if not hujjat:
        return ''
    parts = [hujjat]
    if sana:
        parts.append(f"{sana} йил")
    if organ:
        parts.append(f"{organ}дан берилган")
    return ', '.join(parts) if len(parts) > 1 else parts[0]


def convert_docx_to_pdf(docx_path, pdf_path=None, delete_docx=False):
    """
    .docx faylni .pdf ga aylantiradi. Windows'da MS Word o'rnatilgan bo'lishi
    shart (docx2pdf shu orqali ishlaydi). Agar Word bo'lmasa, xato chiqadi.
    """
    # PyInstaller --windowed rejimida sys.stdout/stderr None bo'lishi mumkin —
    # docx2pdf shunga yozishga urinib xato beradi, shu sabab himoya qo'yamiz.
    if sys.stdout is None:
        sys.stdout = open(os.devnull, 'w')
    if sys.stderr is None:
        sys.stderr = open(os.devnull, 'w')

    if pdf_path is None:
        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
    try:
        from docx2pdf import convert
    except ImportError:
        raise RuntimeError(
            "PDF yaratish uchun 'docx2pdf' kutubxonasi o'rnatilmagan. "
            "requirements.txt orqali o'rnating: pip install docx2pdf pywin32"
        )
    try:
        convert(docx_path, pdf_path)
    except Exception as e:
        raise RuntimeError(
            f"PDF'ga aylantirishda xato: {e}\n"
            "Bu funksiya faqat Windows'da, Microsoft Word o'rnatilgan bo'lsa ishlaydi."
        )
    if delete_docx and os.path.exists(docx_path):
        os.remove(docx_path)
    return pdf_path


def generate_yigma_jild_titul(output_path, xat, portfel_row, mijoz, settings):
    """
    Ijro harakati (MIB) yig'ma jildi uchun titul (muqova) hujjatini,
    yuklab bo'ladigan shablon (YIGMA_JILD_TITUL_TEMPLATE_PATH) asosida
    yaratadi. Bu qog'oz holatidagi ish papkasining birinchi varag'i bo'lib,
    ish bo'yicha barcha asosiy ma'lumotlarni bir joyda ko'rsatadi.
    """
    doc = Document(YIGMA_JILD_TITUL_TEMPLATE_PATH)

    mijoz_ism = (mijoz.get('ism') if mijoz else None) or xat.get('mijoz_nomi', '') or \
        portfel_row.get('mijoz_nomi', '')
    mijoz_ism = util.mijoz_ism_hujjat_uchun(mijoz_ism, xat.get('mijoz_turi'))
    mijoz_manzil = (mijoz.get('manzil') if mijoz else '') or ''
    hujjat_raqami = (mijoz.get('hujjat_raqami') if mijoz else '') or ''
    jami_qarz = (portfel_row.get('asosiy_qarz', 0) or 0) + (portfel_row.get('foiz_qarz', 0) or 0) + \
        (portfel_row.get('jarima', 0) or 0)

    mapping = {
        'BANK_NOMI': settings.get('bank_nomi', ''),
        'FILIAL_NOMI': settings.get('filial_nomi', ''),
        'MIB_ISH_RAQAMI': xat.get('mib_ish_raqami', '') or '—',
        'QARZDOR_ISM': mijoz_ism,
        'MIJOZ_TURI': {'yuridik': "Yuridik shaxs", 'yatt': "Yakka tartibdagi tadbirkor (YaTT)"}.get(
            xat.get('mijoz_turi'), "Jismoniy shaxs"),
        'MIJOZ_MANZIL': mijoz_manzil or '—',
        'HUJJAT_RAQAMI': hujjat_raqami or '—',
        'ANKETA_RAQAM': xat.get('anketa_raqami', '') or '—',
        'JAMI_QARZ': f"{jami_qarz:,.0f}".replace(',', ' '),
        'XAT_SANASI': _fmt_sana_iso(xat.get('yaratilgan_sana')) or '—',
        'XAT_YUBORILGAN_SANA': _fmt_sana_iso(xat.get('yuborilgan_sana')) or '—',
        'DAVO_ARIZA_SANA': _fmt_sana_iso(xat.get('davo_ariza_sana')) or '—',
        'PALATADAN_QAYTGAN_SANA': xat.get('davo_ariza_imzo_sana', '') or '—',
        'SUDGA_TOPSHIRILGAN_SANA': xat.get('sud_topshirilgan_sana', '') or '—',
        'SUD_ISH_RAQAMI': xat.get('sud_ish_raqami', '') or '—',
        'MIB_OTKAZILGAN_SANA': xat.get('mib_otkazilgan_sana', '') or '—',
        'JILD_OCHILGAN_SANA': datetime.date.today().strftime('%d.%m.%Y'),
    }

    _replace_everywhere(doc, mapping)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _fmt_sana_iso(sana_str):
    if not sana_str:
        return ''
    try:
        return datetime.datetime.fromisoformat(sana_str).strftime('%d.%m.%Y')
    except Exception:
        return sana_str


def generate_sugurta_xabarnoma(output_path, vafot_row, portfel_row, mijoz, settings):
    """
    Vafot etgan mijozning kreditini qoplash so'rovi bilan sug'urta
    kompaniyasiga yuboriladigan xabarnoma hujjatini tayyorlaydi.
    """
    doc = Document(SUGURTA_XABARNOMA_TEMPLATE_PATH)

    jami_qarz = (portfel_row.get('asosiy_qarz', 0) or 0) + (portfel_row.get('foiz_qarz', 0) or 0) + \
        (portfel_row.get('jarima', 0) or 0)
    mijoz_ism = (mijoz.get('ism') if mijoz else None) or vafot_row.get('mijoz_nomi', '') or \
        portfel_row.get('mijoz_nomi', '')

    mapping = {
        'BANK_NOMI': settings.get('bank_nomi', ''),
        'BANK_QISQA_NOMI': settings.get('bank_qisqa_nomi', ''),
        'FILIAL_NOMI': settings.get('filial_nomi', ''),
        'RAHBAR_ISM': settings.get('rahbar_ism', ''),

        'MIJOZ_ISM': mijoz_ism,
        'MIJOZ_PINFL': portfel_row.get('pinfl', '') or (mijoz.get('hujjat_raqami') if mijoz else '') or '',
        'ANKETA_RAQAM': vafot_row.get('anketa_raqami', '') or portfel_row.get('anketa_raqami', ''),
        'SHARTNOMA_SANA': portfel_row.get('shartnoma_sanasi', '') or '',
        'KREDIT_SUMMA': _fmt_summa(portfel_row.get('jami_berilgan_summa') or portfel_row.get('ead', 0)),
        'KREDIT_TUGASH_SANASI': vafot_row.get('kredit_tugash_sanasi', '') or
                                 portfel_row.get('shartnoma_tugash_sanasi', '') or '',
        'JAMI_QARZ': _fmt_summa(jami_qarz),

        'VAFOT_SANASI': vafot_row.get('vafot_sanasi', '') or '',
        'SUGURTA_KOMPANIYA': vafot_row.get('sugurta_kompaniya', '') or '',
        'SUGURTA_POLIS_RAQAM': vafot_row.get('sugurta_polis_raqam', '') or '',
        'XABARNOMA_SANA': datetime.date.today().strftime('%d.%m.%Y'),
    }

    _replace_everywhere(doc, mapping)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _bar_chart_png(labels, values, title, color='#1E2761', fmt_short=True):
    """Oddiy gorizontal bar-chart PNG faylini vaqtinchalik joyga chizib beradi."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import tempfile

    fig, ax = plt.subplots(figsize=(7.2, max(2.2, 0.5 * len(labels) + 0.8)))
    y_pos = range(len(labels))
    ax.barh(list(y_pos), values, color=color)
    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()
    ax.set_title(title, fontsize=11, fontweight='bold', loc='left')
    for spine in ('top', 'right'):
        ax.spines[spine].set_visible(False)

    def _short(v):
        if fmt_short:
            if v >= 1e9:
                return f"{v/1e9:.1f} mlrd"
            if v >= 1e6:
                return f"{v/1e6:.1f} mln"
        return f"{v:,.0f}".replace(',', ' ')

    for i, v in enumerate(values):
        ax.text(v, i, ' ' + _short(v), va='center', fontsize=8.5)
    ax.set_xticks([])
    fig.tight_layout()
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    fig.savefig(tmp.name, dpi=150)
    plt.close(fig)
    return tmp.name


def generate_tahlil_hisoboti(output_path, tahlil, settings):
    """"Tahlil" bo'limi uchun to'liq portfel tahlili hisobotini (Word) tayyorlaydi."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(settings.get('bank_nomi', ''))
    r.bold = True
    r.font.size = Pt(13)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{settings.get('filial_nomi','')} filiali")
    r2.font.size = Pt(11)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("PORTFEL TAHLILI HISOBOTI")
    tr.bold = True
    tr.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(datetime.date.today().strftime('%d.%m.%Y'))
    sr.font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading('Umumiy ko\'rsatkichlar', level=2)
    table = doc.add_table(rows=0, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.add_row().cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Ko\'rsatkich', 'Soni', 'EAD summasi (so\'m)'
    rows_data = [
        ('Jami portfel', tahlil['jami_soni'], tahlil['jami_ead']),
        ('Jismoniy shaxslar', tahlil['jismoniy']['soni'], tahlil['jismoniy']['ead']),
        ('Yuridik shaxslar', tahlil['yuridik']['soni'], tahlil['yuridik']['ead']),
    ]
    for label, soni, ead in rows_data:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{soni:,}".replace(',', ' ')
        row[2].text = f"{ead:,.0f}".replace(',', ' ')

    doc.add_paragraph()
    doc.add_heading('Stage bo\'yicha taqsimot', level=2)
    table2 = doc.add_table(rows=0, cols=4)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.add_row().cells
    hdr2[0].text, hdr2[1].text, hdr2[2].text, hdr2[3].text = 'Stage', 'Soni', 'EAD summasi', 'Ulush (%)'
    for st in tahlil['stage']:
        row = table2.add_row().cells
        row[0].text = f"Stage {st['stage']}"
        row[1].text = f"{st['soni']:,}".replace(',', ' ')
        row[2].text = f"{st['ead']:,.0f}".replace(',', ' ')
        row[3].text = f"{st['ulush']}%"

    doc.add_paragraph()
    doc.add_heading('Tarmoq (soha) kesimida', level=2)
    tarmoq_top = tahlil['tarmoq'][:10]
    labels = [t['tarmoq'][:35] for t in tarmoq_top]
    values = [t['ead'] for t in tarmoq_top]
    if labels:
        try:
            chart_path = _bar_chart_png(labels, values, "Tarmoq bo'yicha EAD summasi")
            doc.add_picture(chart_path, width=Inches(6.3))
        except ImportError:
            # matplotlib mavjud bo'lmasa, hisobot baribir jadval bilan davom etadi
            note = doc.add_paragraph()
            note.add_run("(Grafik ko'rinishi mavjud emas — matplotlib kutubxonasi topilmadi. "
                          "Jadval quyida ko'rsatilgan.)").italic = True

    table3 = doc.add_table(rows=0, cols=3)
    table3.style = 'Light Grid Accent 1'
    hdr3 = table3.add_row().cells
    hdr3[0].text, hdr3[1].text, hdr3[2].text = 'Tarmoq', 'Soni', 'EAD summasi'
    for t in tahlil['tarmoq']:
        row = table3.add_row().cells
        row[0].text = t['tarmoq']
        row[1].text = f"{t['soni']:,}".replace(',', ' ')
        row[2].text = f"{t['ead']:,.0f}".replace(',', ' ')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_reja_hisoboti(output_path, reja, tarmoq_reja, settings):
    """"Reja Grafik" bo'limi uchun kunlik ish rejasi hisobotini (Word) tayyorlaydi."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(settings.get('bank_nomi', ''))
    r.bold = True
    r.font.size = Pt(13)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("KUNLIK ISH REJASI HISOBOTI")
    tr.bold = True
    tr.font.size = Pt(18)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(datetime.date.today().strftime('%d.%m.%Y'))
    sr.font.size = Pt(10)
    doc.add_paragraph()

    doc.add_heading("Bugungi reja / bajarilish", level=2)
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.add_row().cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Ish turi', 'Kerak', 'Bajarildi', 'Qolib ketyapti'
    for turk in reja['turkumlar']:
        row = table.add_row().cells
        row[0].text = turk['nomi']
        row[1].text = str(turk['reja'])
        row[2].text = str(turk['bajarildi'])
        row[3].text = str(turk['qoldi'])
    doc.add_paragraph(f"Umumiy bajarilish: {reja['foiz']}%")

    doc.add_paragraph()
    doc.add_heading("Tarmoq kesimida bajarilgan ishlar", level=2)
    table2 = doc.add_table(rows=0, cols=6)
    table2.style = 'Light Grid Accent 1'
    hdr2 = table2.add_row().cells
    for i, h in enumerate(['Tarmoq', 'Jami xat', 'Yuborilgan', 'Davo ariza', 'Sudga', 'MIBga']):
        hdr2[i].text = h
    for t in tarmoq_reja:
        row = table2.add_row().cells
        row[0].text = t['tarmoq']
        row[1].text = str(t['xat_soni'])
        row[2].text = str(t['yuborilgan_soni'])
        row[3].text = str(t['davo_soni'])
        row[4].text = str(t['sud_soni'])
        row[5].text = str(t['mib_soni'])

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
